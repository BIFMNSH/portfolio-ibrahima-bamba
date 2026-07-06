from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "assets" / "cv-ibrahima-bamba.docx"

GREEN = "24583E"
INK = "1A1F1C"
SOFT = "4A5249"


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_bottom_border(paragraph, color=GREEN, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepNext"))


def add_section_heading(document, text):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2.5)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    set_bottom_border(paragraph)
    return paragraph


def add_body(document, text, *, bold_prefix=None, italic=False, space_after=1.5):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.03
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(INK)
        run.italic = italic
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.36)
    paragraph.paragraph_format.first_line_indent = Cm(-0.24)
    paragraph.paragraph_format.space_after = Pt(1.1)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run(f"- {text}")
    run.font.name = "Arial"
    run.font.size = Pt(8.7)
    run.font.color.rgb = RGBColor.from_string(INK)
    return paragraph


def add_role(document, dates, role, organisation):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(15.2)
    table.columns[1].width = Cm(3.1)
    table.rows[0].cells[0].width = Cm(15.2)
    table.rows[0].cells[1].width = Cm(3.1)
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    role_run = table.cell(0, 0).paragraphs[0].add_run(role)
    role_run.bold = True
    role_run.font.name = "Arial"
    role_run.font.size = Pt(9.4)
    role_run.font.color.rgb = RGBColor.from_string(INK)

    date_run = table.cell(0, 1).paragraphs[0].add_run(dates)
    date_run.bold = True
    date_run.font.name = "Arial"
    date_run.font.size = Pt(8.8)
    date_run.font.color.rgb = RGBColor.from_string(GREEN)

    organisation_paragraph = document.add_paragraph()
    organisation_paragraph.paragraph_format.space_after = Pt(0.5)
    organisation_run = organisation_paragraph.add_run(organisation)
    organisation_run.italic = True
    organisation_run.font.name = "Arial"
    organisation_run.font.size = Pt(8.5)
    organisation_run.font.color.rgb = RGBColor.from_string(SOFT)
    keep_with_next(organisation_paragraph)


def build_cv():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.9)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(1.5)

    heading = styles["Heading 1"]
    heading.font.name = "Arial"
    heading.font.size = Pt(11)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(GREEN)
    heading.paragraph_format.keep_with_next = True

    properties = document.core_properties
    properties.title = "CV - Ibrahima Fa Bamba"
    properties.subject = "Chef de projet rénovation énergétique junior"
    properties.author = "Ibrahima Fa Bamba"
    properties.keywords = "rénovation énergétique, copropriété, audit énergétique, coordination, aides"

    header = document.add_table(rows=1, cols=1)
    header.autofit = False
    header.columns[0].width = Cm(18.3)
    cell = header.cell(0, 0)
    set_cell_margins(cell, top=90, start=150, bottom=95, end=150)
    shade_cell(cell, GREEN)

    name = cell.paragraphs[0]
    name.paragraph_format.space_after = Pt(0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run("IBRAHIMA FA BAMBA")
    name_run.font.name = "Arial"
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(255, 255, 255)

    role = cell.add_paragraph()
    role.paragraph_format.space_after = Pt(2)
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = role.add_run("Chef de projet rénovation énergétique junior")
    role_run.font.name = "Arial"
    role_run.font.size = Pt(10.5)
    role_run.font.bold = True
    role_run.font.color.rgb = RGBColor(255, 255, 255)

    contact = cell.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    contact_run = contact.add_run(
        "Freneuse (78) | Mobilité Île-de-France | Permis B | 06 95 42 36 28\n"
        "bamba.bif@gmail.com | linkedin.com/in/fa-ibrahima-bamba"
    )
    contact_run.font.name = "Arial"
    contact_run.font.size = Pt(8)
    contact_run.font.color.rgb = RGBColor(255, 255, 255)

    add_section_heading(document, "Profil")
    add_body(
        document,
        "Scientifique avec plus de quatre ans d’expérience en recherche, analyse de données, conduite de projets et vulgarisation. En reconversion vers la rénovation énergétique, je mobilise cette rigueur pour analyser des dossiers, comparer des scénarios et rendre les décisions technico-financières accessibles aux copropriétaires et aux équipes projet.",
    )

    add_section_heading(document, "Compétences clés")
    add_body(
        document,
        "Audit énergétique et déperditions | Scénarios de rénovation globale | Coûts, gains, aides et reste à charge | Coordination des acteurs et préparation chantier | Conformité, réception et DOE | Vulgarisation technique et supports de décision",
    )

    add_section_heading(document, "Projets de mise en pratique - OpenClassrooms")
    add_body(
        document,
        "Parcours certifiant Chef de projet en rénovation énergétique - RNCP niveau 5 (bac +2). Travaux réalisés à partir de données pédagogiques.",
        italic=True,
        space_after=1,
    )
    add_bullet(document, "Réaliser des états des lieux énergétiques à partir de données client, factures, photos, Visuary et CAP Renov+.")
    add_bullet(document, "Comparer des scénarios pour maisons individuelles et copropriétés en intégrant coûts, gains, aides et reste à charge.")
    add_bullet(document, "Construire une trajectoire de rénovation pour une maison classée passoire thermique.")
    add_bullet(document, "Préparer le lancement d’un chantier : devis, qualifications RGE, autorisations, aides, planning et interfaces.")
    add_bullet(document, "Structurer le contrôle des travaux, la réception et la communication d’un projet de rénovation globale en copropriété.")

    add_section_heading(document, "Expériences professionnelles")
    add_role(
        document,
        "08/2024 - Présent",
        "Enseignant de physique-chimie",
        "Éducation nationale - Académie de Versailles - CDD",
    )
    add_bullet(document, "Concevoir et animer des séquences sur l’énergie, la matière et l’environnement.")
    add_bullet(document, "Vulgariser des notions scientifiques complexes auprès d’élèves non spécialistes.")
    add_bullet(document, "Gérer un groupe et collaborer au sein d’une équipe pédagogique.")

    add_role(
        document,
        "09/2019 - 06/2023",
        "Chargé de projets scientifiques en chimie et électrochimie",
        "Institut Lavoisier de Versailles - Université Paris-Saclay - CDD",
    )
    add_bullet(document, "Conduire des projets de recherche sur des matériaux pour l’énergie : batteries et électrochimie.")
    add_bullet(document, "Concevoir, planifier et suivre des expérimentations en laboratoire.")
    add_bullet(document, "Analyser les données et rédiger des rapports techniques et contributions scientifiques.")

    add_role(
        document,
        "01/2017 - 07/2017",
        "Assistant de recherche en catalyse et spectroscopie",
        "Institut Lavoisier de Versailles - Université Paris-Saclay - CDD",
    )
    add_bullet(document, "Réaliser des tests catalytiques et des analyses spectroscopiques.")
    add_bullet(document, "Optimiser des protocoles expérimentaux et structurer la présentation des résultats.")

    add_section_heading(document, "Formation")
    add_body(document, "Chef de projet en rénovation énergétique - OpenClassrooms - Parcours certifiant RNCP niveau 5 (bac +2)")
    add_body(document, "Doctorat en chimie inorganique - Université Paris-Saclay - 2019 à 2023")
    add_body(document, "Master chimie et physicochimie - Université Paris-Saclay - 2017")

    add_section_heading(document, "Outils et langues")
    add_body(
        document,
        "CAP Renov+ | Visuary | Pléiades | France Rénov’ | Excel et Google Sheets | Géoportail de l’urbanisme | Géorisques",
    )
    add_body(document, "Français : langue maternelle | Malinké : langue maternelle | Anglais : niveau technique")

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_cv()
